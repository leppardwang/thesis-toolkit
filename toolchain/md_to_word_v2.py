#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
博士论文MD→Word转换器 v2
支持：标题层级、表格、引用块、加粗、分页符
"""

import re
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 使用命令行参数，默认为当前目录
BASE = sys.argv[1] if len(sys.argv) > 1 else "."
CHAPTERS = [
    "01-第一章_绪论.md",
    "02-第二章_理论基础与文献综述.md",
    "03-第三章_部署模式安全特性分析.md",
    "04-第四章_DS-CRIF 风险识别框架.md",
    "05-第五章_系统实证验证与案例分析.md",
    "06-第六章_结论与展望.md",
]
OUT = sys.argv[2] if len(sys.argv) > 2 else "博士论文_完整版.docx"

def add_formatted_text(para, text):
    """处理加粗(**)、斜体(*)"""
    parts = re.split(r'(\*{1,2}.*?\*{1,2})', text)
    for p in parts:
        if p.startswith('**') and p.endswith('**') and len(p) > 4:
            run = para.add_run(p[2:-2])
            run.bold = True
        elif p.startswith('*') and p.endswith('*') and not p.startswith('**') and len(p) > 2:
            run = para.add_run(p[1:-1])
            run.italic = True
        else:
            if p:
                para.add_run(p)

def process_chapter(doc, filepath):
    print(f"  读取: {os.path.basename(filepath)}")
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = [l.rstrip('\n') for l in f.readlines()]

    i = 0
    buf = []  # 表格缓冲

    while i < len(lines):
        line = lines[i].strip()

        # 跳过---分隔线
        if line.startswith('---'):
            i += 1
            continue

        # 表格检测
        if line.startswith('|') and '|' in line[1:]:
            buf = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                buf.append(lines[i].strip())
                i += 1
            # 渲染表格
            render_table(doc, buf)
            buf = []
            continue

        # 标题
        h = re.match(r'^(#{1,4})\s+(.*)', line)
        if h:
            level = len(h.group(1))
            text = h.group(2).strip()
            text = re.sub(r'^\d+\.\s*', '', text)  # 去掉"第一章"等前缀
            try:
                para = doc.add_heading(text, level=level)
            except Exception:
                para = doc.add_paragraph(text)
                para.style = f'Heading {level}'
            i += 1
            continue

        # 引用块
        if line.startswith('>'):
            text = line.lstrip('> ').strip()
            para = doc.add_paragraph()
            para.style = 'Intense Quote'
            add_formatted_text(para, text)
            i += 1
            continue

        # 代码块（跳过）
        if line.startswith('```'):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            i += 1
            continue

        # 空行→段落分隔
        if not line:
            i += 1
            continue

        # 普通段落
        para = doc.add_paragraph()
        add_formatted_text(para, line)
        i += 1

    doc.add_page_break()

def render_table(doc, buf):
    """渲染表格（简单版）"""
    rows = []
    for line in buf:
        cells = [c.strip() for c in line.split('|')]
        # 去掉首尾空字符串
        while cells and cells[0] == '':
            cells.pop(0)
        while cells and cells[-1] == '':
            cells.pop()
        # 跳过分隔行
        if all(re.match(r'^-+$', c.strip()) for c in cells):
            continue
        rows.append(cells)

    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci < n_cols:
                cell = table.rows[ri].cells[ci]
                cell.text = cell_text
                if ri == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

def main():
    doc = Document()
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    # 尝试设置中文字体
    try:
        style.element.rPr.rFonts.set('w:eastAsia', '宋体')
    except Exception:
        pass

    print("开始转换博士论文（共6章）...")
    for ch in CHAPTERS:
        fp = os.path.join(BASE, ch)
        if os.path.exists(fp):
            process_chapter(doc, fp)
        else:
            print(f"  ⚠️ 文件不存在: {fp}")

    doc.save(OUT)
    print(f"\n✅ 完成! 输出: {OUT}")
    print(f"   文件大小: {os.path.getsize(OUT) / 1024:.1f} KB")

if __name__ == '__main__':
    main()
