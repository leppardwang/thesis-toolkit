#!/usr/bin/env python3
"""TeX章节 → Word文档 转换脚本（含图表插入）"""
import os, re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 配置
SRC = r"C:\Users\Administrator\Documents\短暂同步\01-博士论文主干材料\准备2025.6\4 日志行为模型\【前置的联邦学习】medical-federated-learning0418博论设计备胎\thesis博论\chapters"
FIG_DIR = r"C:\Users\Administrator\Documents\短暂同步\01-博士论文主干材料\准备2025.6\4 日志行为模型\【前置的联邦学习】medical-federated-learning0418博论设计备胎\thesis博论\figures"
NEW_FIG_DIR = r"C:\Users\Administrator\WorkBuddy\Claw\figures"
OUT = r"C:\Users\Administrator\WorkBuddy\Claw\thesis_output"
os.makedirs(OUT, exist_ok=True)

def strip_tex(content):
    """去掉tex命令，只保留可见文本"""
    content = re.sub(r'\\textbf\{(.+?)\}', r'\1', content)
    content = re.sub(r'\\cite\{[^}]*\}', '', content)
    content = re.sub(r'\\label\{[^}]*\}', '', content)
    content = re.sub(r'\\ref\{[^}]*\}', '[引用]', content)
    content = re.sub(r'\\%', '%', content)
    content = re.sub(r'\\$', '', content)
    content = re.sub(r'\$([^$]+)\$', lambda m: m.group(1), content)
    content = re.sub(r'\\begin\{equation\}.*?\\end\{equation\}', '[公式]', content, flags=re.DOTALL)
    content = re.sub(r'\\begin\{table\}.*?\\end\{table\}', '[表格]', content, flags=re.DOTALL)
    content = re.sub(r'\\begin\{algorithm\}.*?\\end\{algorithm\}', '[算法伪代码]', content, flags=re.DOTALL)
    content = re.sub(r'\\begin\{figure\}.*?\\end\{figure\}', '', content, flags=re.DOTALL)
    content = re.sub(r'\\[a-zA-Z]+(\{[^}]*\})?', '', content)
    return content.strip()

def find_image(path, filename):
    """在多个目录查找图片文件"""
    for base in [path, FIG_DIR, NEW_FIG_DIR]:
        for ext in ['.png', '.jpg', '.pdf']:
            full = os.path.join(base, filename + ext)
            if os.path.exists(full):
                return full
    return None

def tex_to_docx(tex_path, doc):
    """单个tex文件转Word段落（含图表插入）"""
    with open(tex_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 提取标题
    title_m = re.search(r'\\chapter\{(.+?)\}', content)
    if title_m:
        doc.add_heading(title_m.group(1), level=1)
    
    # 拆分段落：先按行处理
    for line in content.split('\n'):
        line = line.strip()
        if not line or line.startswith('%'):
            continue
        line = re.sub(r'(?<!\\)%.*$', '', line).strip()
        if not line:
            continue
        
        # 标题
        sec_m = re.search(r'\\section\{(.+?)\}', line)
        if sec_m:
            doc.add_heading(sec_m.group(1), level=2)
            continue
        sub_m = re.search(r'\\subsection\{(.+?)\}', line)
        if sub_m:
            doc.add_heading(sub_m.group(1), level=3)
            continue
        subsub_m = re.search(r'\\subsubsection\{(.+?)\}', line)
        if subsub_m:
            doc.add_heading(subsub_m.group(1), level=4)
            continue
        
        # 图片插入
        img_m = re.search(r'\\includegraphics(?:\[.*?\])?\{(.+?)\}', line)
        if img_m:
            img_name = os.path.splitext(os.path.basename(img_m.group(1)))[0]
            # 找图片标题
            img_path = find_image(os.path.dirname(tex_path), img_name)
            if img_path and os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(4.5))
                continue
        
        # 普通段落
        text = strip_tex(line)
        text = re.sub(r'\s+', ' ', text).strip()
        if len(text) > 15:
            p = doc.add_paragraph(text)
            p.paragraph_format.first_line_indent = Pt(24)

# ========== 主程序 ==========
doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(12)
style.font.name = 'Times New Roman'
chapters = sorted([f for f in os.listdir(SRC) if f.endswith('.tex')])

for ch in chapters:
    tex_path = os.path.join(SRC, ch)
    print(f'转换: {ch}')
    tex_to_docx(tex_path, doc)
    doc.add_page_break()

out_path = os.path.join(OUT, '联邦学习博论_完整稿(含图).docx')
doc.save(out_path)

size_kb = os.path.getsize(out_path) / 1024
print(f'\n✅ 完成! 共 {len(chapters)} 章 → {out_path}')
print(f'   文件大小: {size_kb:.0f} KB = {size_kb/1024:.1f} MB')
